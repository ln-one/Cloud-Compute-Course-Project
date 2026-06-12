from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
COVER_TEMPLATE = ROOT / "templates" / "cover-template-midterm.docx"
OUTPUT = ROOT / "sample-output.docx"


def insert_paragraph_after(paragraph, text: str = "", style_name: str | None = None):
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style_name:
        new_paragraph.style = style_name
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def set_cover(doc: Document, title: str) -> None:
    cover_values = {
        11: ("题目：", title),
        12: ("学院：", "计算机与人工智能学院    "),
        13: ("教师：", "叶家鸣            "),
        14: ("学号姓名：", "张春冉 2023112573   "),
    }
    for paragraph_index, (label, value) in cover_values.items():
        paragraph = doc.paragraphs[paragraph_index]
        if len(paragraph.runs) >= 2:
            paragraph.runs[0].text = label
            paragraph.runs[1].text = value
            for run in paragraph.runs[2:]:
                run.text = ""
        else:
            paragraph.text = label + value


def trim_body_keep_cover(doc: Document) -> None:
    for block in list(doc._body._body)[19:]:
        if block.tag.endswith("sectPr"):
            continue
        parent = block.getparent()
        if parent is not None:
            parent.remove(block)


def set_run_font(run, size: float, bold: bool = False, font_name: str = "Times New Roman"):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold


def add_heading1(anchor, text: str):
    paragraph = insert_paragraph_after(anchor, "", "Normal")
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = Pt(24)
    run = paragraph.add_run(text)
    set_run_font(run, 14, True)
    return paragraph


def add_heading2(anchor, text: str):
    paragraph = insert_paragraph_after(anchor, "", "Normal")
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = Pt(22)
    run = paragraph.add_run(text)
    set_run_font(run, 12.5, True)
    return paragraph


def add_para(anchor, text: str):
    paragraph = insert_paragraph_after(anchor, "", "Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(24)
    run = paragraph.add_run(text)
    set_run_font(run, 12)
    return paragraph


def add_code(anchor, code: str):
    paragraph = insert_paragraph_after(anchor, "", "Normal")
    paragraph.paragraph_format.line_spacing = Pt(13)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(code.rstrip())
    set_run_font(run, 8.5, False, "Consolas")
    return paragraph


def add_image(anchor, image_path: Path, caption: str, width: float = 5.7):
    paragraph = insert_paragraph_after(anchor, "", "Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))

    caption_para = insert_paragraph_after(paragraph, "", "Normal")
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_para.paragraph_format.space_before = Pt(3)
    caption_para.paragraph_format.line_spacing = Pt(20)
    caption_run = caption_para.add_run(caption)
    set_run_font(caption_run, 10.5)
    return caption_para


def main() -> None:
    doc = Document(COVER_TEMPLATE)
    set_cover(doc, "软件设计模式实验-模式名称")
    trim_body_keep_cover(doc)

    anchor = doc.paragraphs[-1]
    anchor = add_heading1(anchor, "一、模式简述")
    anchor = add_para(anchor, "在这里写模式定义、适用场景和本实验中的应用。")
    anchor = add_heading1(anchor, "二、系统设计与实现")
    anchor = add_heading2(anchor, "2.1 需求分析")
    anchor = add_para(anchor, "在这里写实验需求和设计思路。")
    anchor = add_heading1(anchor, "三、主要代码")
    anchor = add_code(anchor, "public class Example {\n    // code here\n}")
    anchor = add_heading1(anchor, "四、运行截图")
    anchor = add_para(anchor, "在这里插入截图和说明。")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
